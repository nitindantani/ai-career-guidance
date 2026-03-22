from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

# Create Word document for Lean Canvas
doc = Document()

# Title
title = doc.add_heading("Lean Canvas: AI-Based Career Guidance System", level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Define Lean Canvas sections
lean_canvas = {
    "1. Problem": [
        "Lack of proper career guidance for students.",
        "Confusion due to too many career options.",
        "Limited access to personalized mentoring."
    ],
    "2. Customer Segments": [
        "High school and college students",
        "Career counselors and mentors",
        "Educational institutions"
    ],
    "3. Unique Value Proposition": [
        "AI-powered personalized career guidance.",
        "Instant predictions and GPT-based chatbot support.",
        "Accessible platform with a simple interface."
    ],
    "4. Solution": [
        "Machine Learning-based career prediction model.",
        "Interactive UI with real-time autocomplete.",
        "Fallback support via GPT chatbot."
    ],
    "5. Channels": [
        "Web application (desktop and mobile)",
        "Institutional platforms",
        "GitHub and Render deployment"
    ],
    "6. Revenue Streams": [
        "Freemium model for individual users.",
        "Subscription for institutions (future scope)."
    ],
    "7. Cost Structure": [
        "API usage fees (OpenAI)",
        "Hosting and domain costs",
        "Development and maintenance"
    ],
    "8. Key Metrics": [
        "User engagement and traffic",
        "Prediction accuracy",
        "Chatbot interaction rate"
    ],
    "9. Unfair Advantage": [
        "Integration of GPT-4 for fallback support.",
        "Custom-trained model on domain-specific data.",
        "User-friendly design focused on student needs."
    ]
}

# Add each section to the document
for section, points in lean_canvas.items():
    doc.add_heading(section, level=2)
    for point in points:
        doc.add_paragraph(f"- {point}", style='List Bullet')

# Save the Word document
docx_path = r"C:\Users\NITIN\Documents/Lean_Canvas_AI_Career_Guidance.docx"
pdf_path = r"C:\Users\NITIN\Documents/Lean_Canvas_AI_Career_Guidance.pdf"
doc.save(docx_path)

# Convert to PDF
convert(docx_path, pdf_path)

pdf_path
