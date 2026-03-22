from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# Create a new Word document
doc = Document()

# Define title style
def add_heading(text, level=1):
    doc.add_heading(text, level=level)

# Define paragraph with normal style
def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    font = run.font
    font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(6)
    return p

# Title Page
doc.add_heading("Final Project Report", 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("AI-Based Career Guidance System", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Student Name: Nitin Dantani").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Institute: Saffrony Institute of Technology").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Department: Computer Science and Engineering").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Date: July 2025").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_page_break()

# Abstract
add_heading("Abstract")
add_paragraph("Choosing a suitable career path is a critical yet challenging decision for students, especially when lacking proper guidance. "
              "This project presents an AI-Based Career Guidance System that leverages machine learning and GPT-based chatbot capabilities to provide personalized career suggestions. "
              "The system processes user inputs like stream, subject interest, technical and soft skills to predict suitable career options and also engages users in intelligent conversation for career queries. "
              "The model is trained using a structured dataset and deployed with a user-friendly web interface.")

# Table of Contents
add_heading("Table of Contents")
contents = [
    "1. Introduction",
    "2. Problem Statement",
    "3. Objectives",
    "4. Target Users",
    "5. Dataset Overview",
    "6. Technology Stack",
    "7. System Architecture",
    "8. Key Features",
    "9. Machine Learning Model",
    "10. Implementation Screenshots",
    "11. Challenges Faced",
    "12. Future Scope",
    "13. Conclusion",
    "14. GitHub & Deployment Links",
    "15. References",
    "16. Alignment with Global Goals"
]
for item in contents:
    add_paragraph(item)

doc.add_page_break()

# Sections
sections = {
    "1. Introduction": "Career decisions are often made under peer pressure or without proper insight into one’s skills and interests. "
                       "This project aims to assist students in making informed career choices using data-driven approaches. "
                       "It combines predictive analytics and AI chat capabilities to provide guidance based on individual preferences and capabilities.",
    
    "2. Problem Statement": "Many students face difficulty selecting a suitable career due to the lack of structured guidance, especially in early academic stages. "
                            "With a wide range of professions available today, it becomes overwhelming to match personal interests with career opportunities. "
                            "The aim of this system is to automate career counseling using AI and ML technologies.",

    "3. Objectives": "\n".join([
        "- Provide automated career suggestions using user preferences",
        "- Utilize ML models trained on real-world data to ensure accuracy",
        "- Offer real-time responses through an AI chatbot",
        "- Ensure user-friendly interaction through autocomplete inputs",
        "- Deploy the system for global accessibility"
    ]),

    "4. Target Users": "\n".join([
        "- High school students",
        "- Undergraduate college students",
        "- Career counselors and mentors",
        "- Educational platforms and institutions"
    ]),

    "5. Dataset Overview": "Name: career_data.csv\nSize: 1000+ records\nFeatures: Stream, Subject_Liked, Skills, Soft_Skill, Preferred_Field, Career_Label\n"
                           "The dataset was cleaned and preprocessed using normalization and encoding techniques. It played a vital role in model training.",

    "6. Technology Stack": "\n".join([
        "- Frontend: HTML, CSS, JavaScript",
        "- Backend: Python, Flask",
        "- Machine Learning: Scikit-learn, pandas",
        "- AI Integration: OpenAI GPT-4 API",
        "- Deployment: Render, GitHub"
    ]),

    "7. System Architecture": "The system architecture consists of user input interface, backend ML model for predictions, GPT chatbot integration for fallback suggestions, and a deployment layer accessible through a browser.",

    "8. Key Features": "\n".join([
        "- Intelligent career prediction form",
        "- Real-time autocomplete using CSV data",
        "- GPT-based AI chatbot for Q&A",
        "- Mobile responsive interface",
        "- Error handling using AI fallback"
    ]),

    "9. Machine Learning Model": "- Model Used: Random Forest Classifier\n- Accuracy Achieved: ~7.5%\n- Preprocessing: Label Encoding\n- Training: 80/20 train-test split\n- Fallback Strategy: GPT-4 suggestion if model prediction fails",

    "10. Implementation Screenshots": "Include screenshots of:\n- Homepage\n- Career prediction form\n- Result display\n- AI chatbot interface\n- Autocomplete fields",

    "11. Challenges Faced": "\n".join([
        "- Low model accuracy due to subjective data",
        "- Validating user input dynamically",
        "- Handling new or unknown input with fallback logic",
        "- Integrating GPT API securely and managing costs"
    ]),

    "12. Future Scope": "\n".join([
        "- Improve accuracy using Deep Learning models",
        "- Add login and user tracking system",
        "- Career suggestions based on uploaded resume",
        "- Localization to Indian languages",
        "- Integration with job and internship platforms"
    ]),

    "13. Conclusion": "This project delivers a practical solution for students who seek career advice based on their academic and personal preferences. "
                      "It showcases the power of AI and ML to solve real-life educational challenges and sets a foundation for future improvements in digital career counseling.",

    "14. GitHub & Deployment Links": "- GitHub: https://github.com/nitindantani/ai_career-guidence\n- Live Demo: https://ai-career-guidance-hdzl.onrender.com",

    "15. References": "\n".join([
        "- https://scikit-learn.org",
        "- https://flask.palletsprojects.com",
        "- https://openai.com/api",
        "- Dataset: Custom-curated CSV based on educational survey",
        "- GitHub and Render documentation"
    ]),

    "16. Alignment with Global Goals": "\n".join([
        "IBM SkillsBuild Relevance:",
        "- Supports IBM’s mission to provide free education and career guidance tools",
        "- Enhances digital literacy and employability through AI-powered guidance",
        "- Helps learners discover career pathways aligned with individual strengths",
        "SDGs (Sustainable Development Goals):",
        "- SDG 4 – Quality Education: Equitable access to planning tools",
        "- SDG 8 – Decent Work and Economic Growth: Encourages job readiness",
        "- SDG 9 – Innovation & Infrastructure: Applies AI in education"
    ])
}

for title, content in sections.items():
    add_heading(title)
    add_paragraph(content)

# Save the document
doc_path = r"C:\Users\NITIN\Documents\AI_Based_Career_Guidance_Report.docx"
doc.save(doc_path)
doc_path
